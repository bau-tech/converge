"""
Office document generators/readers for converge's BIM report pipeline.

Shared IR: a report is (title: str, sections: list[dict]), where each
section is one of:
  {"heading": str, "text": str}                                  narrative
  {"heading": str, "table": {"columns": [...], "rows": [[...]]}}  tabular
  {"heading": str, "images": [{"caption": str, "data": bytes}]}   image gallery
A table cell may itself be {"type": "image", "data": bytes} instead of a
plain value — rendered as a small thumbnail sized to its column/row instead
of text (gather_bcf uses this for each topic's viewpoint screenshot, inline
in the Topics table rather than a separate gallery after it).
reports/generate.py's gather_* functions build this IR once per report
type and hand it to whichever builder below matches the requested
output_format — one IR, three renderers, so a new report type only needs a
data gatherer, not per-format rendering code.
"""

from __future__ import annotations

import io
import os
import re

# Small (160x160, tightly cropped, ~50KB) pre-resized copy of
# public/converge-logo2-transparent.png — kept separate from the 2048x2048/
# 2.8MB source so every generated report doesn't balloon in size. Branding
# is skipped (not an error) if this is ever missing, since a missing logo
# shouldn't block report generation.
_LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "converge-logo.png")


def _sanitize_sheet_name(name: str) -> str:
    """Excel sheet names can't contain \\ / ? * [ ] : and are capped at 31 chars."""
    name = re.sub(r'[\\/?*\[\]:]', "-", name or "Sheet").strip()
    return (name or "Sheet")[:31]


def _is_image_cell(v) -> bool:
    return isinstance(v, dict) and v.get("type") == "image" and bool(v.get("data"))


def _split_title(title: str) -> tuple[str, str | None]:
    """Every gatherer's title is "{Report Name} — {Server-Project-Model-
    Version label}" (or, for gather_changes, two labels joined by " → ").
    That label can be long, so it's rendered as a smaller subtitle under
    the report name rather than crammed into one huge Title-styled line —
    split on the first " — " to separate them."""
    main, sep, rest = title.partition(" — ")
    return (main, rest) if sep else (title, None)


def _add_docx_branding(doc) -> None:
    """Logo + "Converge" in the document's actual header — repeats at the
    top-left of every page, not just the first."""
    if not os.path.isfile(_LOGO_PATH):
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    header = doc.sections[0].header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        para.add_run().add_picture(_LOGO_PATH, height=Inches(0.3))
    except Exception:
        pass
    brand_run = para.add_run("  Converge")
    brand_run.bold = True
    brand_run.font.size = Pt(14)


def build_docx_report(title: str, sections: list[dict]) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    _add_docx_branding(doc)
    main_title, subtitle = _split_title(title)
    doc.add_heading(main_title, level=0)
    if subtitle:
        sub_run = doc.add_paragraph().add_run(subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for section in sections:
        heading = section.get("heading")
        if heading:
            doc.add_heading(heading, level=1)
        table = section.get("table")
        if table:
            columns = table.get("columns", [])
            rows = table.get("rows", [])
            if columns:
                docx_table = doc.add_table(rows=1, cols=len(columns))
                docx_table.style = "Light Grid Accent 1"
                for i, col in enumerate(columns):
                    docx_table.rows[0].cells[i].text = str(col)
                for row in rows:
                    cells = docx_table.add_row().cells
                    for i, val in enumerate(row):
                        if _is_image_cell(val):
                            try:
                                cells[i].paragraphs[0].add_run().add_picture(io.BytesIO(val["data"]), width=Inches(1.1))
                            except Exception:
                                pass
                        else:
                            cells[i].text = "" if val is None else str(val)
        images = section.get("images")
        if images:
            for img in images:
                data_bytes = img.get("data") if isinstance(img, dict) else None
                caption = img.get("caption") if isinstance(img, dict) else None
                if not data_bytes:
                    continue
                try:
                    doc.add_picture(io.BytesIO(data_bytes), width=Inches(5.5))
                except Exception:
                    continue
                if caption:
                    cap_run = doc.add_paragraph().add_run(caption)
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
        if section.get("text"):
            for para in str(section["text"]).split("\n"):
                doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_xlsx_branding(ws) -> None:
    """Logo + "Converge" pinned to the top-left of a sheet. Excel has no
    real repeating "header" concept in normal view (only print headers,
    which aren't visible on screen) — inserting 3 blank rows and anchoring
    the logo at A1 is the closest on-screen equivalent, applied per-sheet
    so it's visible regardless of which tab is active when the workbook
    opens."""
    if not os.path.isfile(_LOGO_PATH):
        return
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Font

    ws.insert_rows(1, 3)
    try:
        img = XLImage(_LOGO_PATH)
        img.width = 36
        img.height = 36
        ws.add_image(img, "A1")
    except Exception:
        pass
    ws["C2"] = "Converge"
    ws["C2"].font = Font(bold=True, size=14)
    ws.row_dimensions[1].height = 30


def build_xlsx_report(title: str, sections: list[dict]) -> bytes:
    """Each table-bearing section becomes its own sheet (named from its
    heading); narrative (text-only) sections are collected onto one
    "Summary" sheet, one line per paragraph."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    wb.remove(wb.active)

    main_title, subtitle = _split_title(title)
    summary_lines = [main_title] + ([subtitle] if subtitle else []) + [""]
    base_len = len(summary_lines)
    for section in sections:
        table = section.get("table")
        if table:
            columns = table.get("columns", [])
            rows = table.get("rows", [])
            sheet_name = _sanitize_sheet_name(section.get("heading") or "Sheet")
            base_name, n = sheet_name, 1
            while sheet_name in wb.sheetnames:
                n += 1
                suffix = f"_{n}"
                sheet_name = f"{base_name[:31 - len(suffix)]}{suffix}"
            ws = wb.create_sheet(sheet_name)
            # Branding MUST go on before any content: it inserts 3 rows at
            # the top, but openpyxl's insert_rows only shifts cell values —
            # it does not re-anchor already-placed images. Adding images
            # first, then branding, left every image anchored 3 rows too
            # high (into the branding band itself) once the insert ran,
            # even though the row's text correctly shifted down beneath it.
            _add_xlsx_branding(ws)
            if columns:
                ws.append(columns)
                for cell in ws[ws.max_row]:
                    cell.font = Font(bold=True)
            image_col = next(
                (i for i in range(len(columns)) if any(_is_image_cell(r[i]) for r in rows if i < len(r))), None,
            )
            for row in rows:
                ws.append(["" if (v is None or _is_image_cell(v)) else v for v in row])
                if image_col is not None and image_col < len(row) and _is_image_cell(row[image_col]):
                    from openpyxl.drawing.image import Image as XLImage
                    try:
                        xl_img = XLImage(io.BytesIO(row[image_col]["data"]))
                        xl_img.width, xl_img.height = 70, 50
                        ws.add_image(xl_img, ws.cell(row=ws.max_row, column=image_col + 1).coordinate)
                        ws.row_dimensions[ws.max_row].height = 42
                    except Exception:
                        pass
            if image_col is not None:
                ws.column_dimensions[ws.cell(row=1, column=image_col + 1).column_letter].width = 12
        elif section.get("images"):
            from openpyxl.drawing.image import Image as XLImage

            sheet_name = _sanitize_sheet_name(section.get("heading") or "Images")
            base_name, n = sheet_name, 1
            while sheet_name in wb.sheetnames:
                n += 1
                suffix = f"_{n}"
                sheet_name = f"{base_name[:31 - len(suffix)]}{suffix}"
            ws = wb.create_sheet(sheet_name)
            _add_xlsx_branding(ws)  # before content — see note above
            for img in section["images"]:
                data_bytes = img.get("data") if isinstance(img, dict) else None
                caption = img.get("caption") if isinstance(img, dict) else None
                if not data_bytes:
                    continue
                if caption:
                    ws.append([caption])
                try:
                    xl_img = XLImage(io.BytesIO(data_bytes))
                    xl_img.width, xl_img.height = 300, 200
                    ws.add_image(xl_img, f"A{ws.max_row + 1}")
                except Exception:
                    continue
                # Blank rows reserve vertical space so the next image/caption
                # doesn't overlap this one — openpyxl anchors images by cell
                # but doesn't grow row heights to fit them.
                for _ in range(11):
                    ws.append([])
        else:
            if section.get("heading"):
                summary_lines.append(f"## {section['heading']}")
            if section.get("text"):
                summary_lines.extend(str(section["text"]).split("\n"))
            summary_lines.append("")

    if len(summary_lines) > base_len or not wb.sheetnames:
        ws = wb.create_sheet("Summary", 0)
        _add_xlsx_branding(ws)  # before content — see note in the table branch above
        title_row = ws.max_row + 1
        for line in summary_lines:
            ws.append([line])
        ws.cell(row=title_row, column=1).font = Font(bold=True, size=14)
        if subtitle:
            ws.cell(row=title_row + 1, column=1).font = Font(italic=True, size=10, color="666666")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _draw_pdf_branding(canvas_obj, doc_template) -> None:
    """Logo + "Converge" top-left, page number bottom-right — both drawn
    directly onto the page canvas since reportlab's page template callback
    is the only way to get content that repeats on every page (Platypus
    flowables only place content inside the margins, not in them)."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    page_width, page_height = doc_template.pagesize
    canvas_obj.saveState()
    if os.path.isfile(_LOGO_PATH):
        logo_h = 1.0 * cm
        x = doc_template.leftMargin
        y = page_height - 1.7 * cm
        try:
            canvas_obj.drawImage(
                _LOGO_PATH, x, y, width=logo_h, height=logo_h,
                mask="auto", preserveAspectRatio=True,
            )
        except Exception:
            pass
        canvas_obj.setFont("Helvetica-Bold", 13)
        canvas_obj.setFillColor(colors.HexColor("#1a1a1a"))
        canvas_obj.drawString(x + logo_h + 0.25 * cm, y + logo_h / 2 - 4, "Converge")
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(colors.HexColor("#888888"))
    canvas_obj.drawRightString(page_width - doc_template.rightMargin, 1.2 * cm, f"Page {canvas_obj.getPageNumber()}")
    canvas_obj.restoreState()


def build_pdf_report(title: str, sections: list[dict]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    # Table cells used to hold raw strings with no colWidths set, so reportlab
    # sized every column to its widest content and just let the table run
    # past the page's right edge for anything with a long free-text column
    # (e.g. BCF topic titles) — Paragraph-wrapped cells + colWidths summing
    # to the page's actual usable width fixes that regardless of content length.
    cell_style = ParagraphStyle("ReportCell", parent=styles["Normal"], fontSize=8, leading=10)
    header_cell_style = ParagraphStyle(
        "ReportCellHeader", parent=cell_style, textColor=colors.white, fontName="Helvetica-Bold",
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=3 * cm, bottomMargin=2 * cm)
    main_title, subtitle = _split_title(title)
    story = [Paragraph(main_title, styles["Title"])]
    if subtitle:
        subtitle_style = styles["Normal"].clone("Subtitle")
        subtitle_style.textColor = colors.HexColor("#666666")
        subtitle_style.fontSize = 10
        story.append(Paragraph(subtitle, subtitle_style))
    story.append(Spacer(1, 0.5 * cm))

    for section in sections:
        if section.get("heading"):
            story.append(Paragraph(section["heading"], styles["Heading2"]))
        table = section.get("table")
        if table:
            columns = table.get("columns", [])
            rows = table.get("rows", [])
            header_row = [Paragraph(str(c), header_cell_style) for c in columns] if columns else None
            # A column that holds thumbnails (e.g. gather_bcf's per-topic
            # viewpoint screenshot) needs a small fixed width instead of an
            # equal share of the page — sized here, then rendered as a real
            # Image flowable per cell rather than Paragraph-wrapped text.
            n_cols = len(columns) if columns else (len(rows[0]) if rows else 0)
            image_col = next(
                (i for i in range(n_cols) if any(_is_image_cell(r[i]) for r in rows if i < len(r))), None,
            )
            image_col_width = 2.4 * cm
            thumb_w, thumb_h = 2.1 * cm, 1.5 * cm

            def _cell(v, style):
                if _is_image_cell(v):
                    try:
                        img = RLImage(io.BytesIO(v["data"]), width=thumb_w, height=thumb_h)
                        img.hAlign = "CENTER"
                        return img
                    except Exception:
                        return Paragraph("", style)
                return Paragraph("" if v is None else str(v), style)

            body_rows = [[_cell(v, cell_style) for v in row] for row in rows]
            data = ([header_row] if header_row else []) + body_rows
            if data:
                if image_col is not None and n_cols > 1:
                    other_width = (doc.width - image_col_width) / (n_cols - 1)
                    col_widths = [other_width] * n_cols
                    col_widths[image_col] = image_col_width
                else:
                    col_widths = [doc.width / n_cols] * n_cols
                t = Table(data, colWidths=col_widths, repeatRows=1 if header_row else 0)
                # Softer header (was near-black navy) + light grid + zebra
                # striping instead of stark white rows — the previous style
                # read as too high-contrast/harsh against the page.
                style_cmds = [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#5a6b87")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d5d5da")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
                body_start = 1 if header_row else 0
                for i in range(body_start, len(data)):
                    if (i - body_start) % 2 == 1:
                        style_cmds.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#f4f5f7")))
                t.setStyle(TableStyle(style_cmds))
                story.append(t)
        images = section.get("images")
        if images:
            max_w = doc.width
            for img in images:
                data_bytes = img.get("data") if isinstance(img, dict) else None
                caption = img.get("caption") if isinstance(img, dict) else None
                if not data_bytes:
                    continue
                try:
                    rl_img = RLImage(io.BytesIO(data_bytes))
                    if rl_img.drawWidth > max_w:
                        scale = max_w / rl_img.drawWidth
                        rl_img.drawWidth *= scale
                        rl_img.drawHeight *= scale
                    story.append(rl_img)
                    if caption:
                        story.append(Paragraph(caption, cell_style))
                    story.append(Spacer(1, 0.3 * cm))
                except Exception:
                    continue
        if section.get("text"):
            for para in str(section["text"]).split("\n"):
                story.append(Paragraph(para or "&nbsp;", styles["Normal"]))
        story.append(Spacer(1, 0.4 * cm))

    doc.build(story, onFirstPage=_draw_pdf_branding, onLaterPages=_draw_pdf_branding)
    return buf.getvalue()


def read_docx_text(data: bytes) -> str:
    """Extract paragraphs and tables as plain text from a .docx file's bytes."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = [para.text for para in doc.paragraphs if para.text.strip()]
    for i, table in enumerate(doc.tables):
        lines.append(f"\n[Table {i + 1}]")
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines) if lines else "(empty document)"


def read_xlsx_text(data: bytes) -> str:
    """Extract sheet names and cell values as plain text from an .xlsx file's bytes."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"\n[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            if any(v is not None for v in row):
                lines.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(lines) if lines else "(empty workbook)"
